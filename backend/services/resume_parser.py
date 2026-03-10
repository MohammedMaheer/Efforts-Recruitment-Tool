import PyPDF2
import docx
import re
import logging
import unicodedata
from typing import Dict, List, Any, Optional
from io import BytesIO

logger = logging.getLogger(__name__)

# Common mojibake sequences (UTF-8 decoded as latin-1/cp1252)
MOJIBAKE_PATTERNS = re.compile(
    r'Ã[\x80-\xbf¡-ÿ]'
    r'|â€[™¢¦œ"\'\-\—]'
    r'|Ã¢â‚¬'
    r'|â\x80[\x90-\x9f]'
    r'|Â[\xa0-\xff]'
    r'|Ã©|Ã¨|Ã¼|Ã¶|Ã±|Ã§'
)

def is_mojibake(text: str, threshold: float = 0.15) -> bool:
    """Detect if text contains excessive mojibake / encoding corruption.
    Returns True if >threshold fraction of chars are garbled."""
    if not text or len(text) < 20:
        return False
    mojibake_chars = sum(len(m) for m in MOJIBAKE_PATTERNS.findall(text))
    ratio = mojibake_chars / len(text)
    return ratio > threshold


# ── Spaced-character corruption detection & repair ──────────────────────
# pdfplumber layout mode sometimes emits single chars separated by spaces/dashes/bullets
# e.g.  "S   A   R   A   V   A   N   A"  or  "S•A•R•A•V•A"
# The separator class includes: whitespace, dashes, dots, underscores, bullets,
# and other common PDF-artifact characters (•, ◦, ▪, ▸, ‣, etc.)
_SEP_CLASS = r'[\s\-._\u2022\u2023\u2043\u25CF\u25CB\u25AA\u25AB\u25B8\u25B9\u2219\u00B7\u2027\u00A0\u2013\u2014\u2015\u007C\u002A\u2018\u2019\u201C\u201D]'
_SPACED_CHAR_RE = re.compile(
    r'(?:^|(?<=\s))'                         # start of string or after whitespace
    r'[A-Za-z0-9]'                            # first char
    r'(?:' + _SEP_CLASS + r'{1,20}[A-Za-z0-9]){4,}'  # 4+ reps of separator+char (up to 20 sep chars between letters)
    r'(?:$|(?=\s))',                           # end or followed by whitespace
    re.MULTILINE
)

def is_spaced_text(text: str, threshold: float = 0.25, min_length: int = 40) -> bool:
    """Return True when a significant fraction of the text consists of
    single characters separated by spaces / dashes / dots / bullets — a hallmark of
    bad pdfplumber layout-mode extraction.
    Use min_length=5 for short strings like names."""
    if not text or len(text) < min_length:
        return False
    # Method 1: Regex matching of spaced-char runs
    matches = _SPACED_CHAR_RE.findall(text)
    matched_chars = sum(len(m) for m in matches)
    ratio = matched_chars / len(text)
    if ratio > threshold:
        return True
    # Method 2: Check ratio of non-alnum to alnum chars (catches bullet-stuffed text)
    alnum = sum(1 for c in text if c.isalnum())
    non_alnum = len(text) - alnum
    if alnum > 10 and non_alnum / len(text) > 0.75:
        # High ratio of non-alnum chars — check if the alnum chars are isolated
        # (i.e., most are surrounded by non-alnum)
        isolated = 0
        for i, c in enumerate(text):
            if c.isalnum():
                prev_alnum = (i > 0 and text[i-1].isalnum())
                next_alnum = (i < len(text)-1 and text[i+1].isalnum())
                if not prev_alnum and not next_alnum:
                    isolated += 1
        if alnum > 0 and isolated / alnum > 0.5:
            return True
    return False


def collapse_spaced_chars(text: str) -> str:
    """Collapse spaced-character runs back into normal words.

    Handles patterns like:
      "S   A   R   A   V   A   N   A"  ->  "SARAVANA"
      "S-A-R-A-V-A-N-A"               ->  "SARAVANA"
      "S•••A•••R•••A"                  ->  "SARA"
    Works per-line so normal sentences are untouched.
    """
    if not text:
        return text

    def _fix_line(line: str) -> str:
        # First try the regex-based approach for well-formed patterns
        def _collapse(m: re.Match) -> str:
            raw = m.group(0)
            letters = re.findall(r'[A-Za-z0-9]', raw)
            return ''.join(letters)
        fixed = _SPACED_CHAR_RE.sub(_collapse, line)

        # If still heavily non-alnum, do aggressive: strip all non-alnum between isolated letters
        alnum = sum(1 for c in fixed if c.isalnum())
        if len(fixed) > 10 and alnum > 0:
            non_alnum_ratio = (len(fixed) - alnum) / len(fixed)
            if non_alnum_ratio > 0.7:
                # Aggressive: extract only alphanumeric chars and spaces
                result = []
                for c in fixed:
                    if c.isalnum() or c in ('\n', ' '):
                        result.append(c)
                fixed = ''.join(result)
                # Collapse single-char sequences: "S A R A V A N A" -> "SARAVANA"
                fixed = re.sub(r'\b(\w)\s+(?=\w\b)', r'\1', fixed)

        return fixed

    lines = text.split('\n')
    return '\n'.join(_fix_line(l) for l in lines)


def text_quality_score(text: str) -> float:
    """Return a 0-1 quality score. Higher = more likely to be readable text.

    Heuristics:
    - Reward common English words
    - Reward word-length variety (real text has short & long words)
    - Penalise single-char "words" (spaced-char corruption artefact)
    - Penalise high ratio of non-alnum characters (bullet-stuffed text)
    """
    if not text or len(text.strip()) < 20:
        return 0.0
    words = text.split()
    if not words:
        return 0.0

    # Fraction of single-char tokens
    single_char_ratio = sum(1 for w in words if len(w) == 1 and w.isalpha()) / len(words)

    # Average word length (good text ~4-7)
    avg_len = sum(len(w) for w in words) / len(words)
    length_score = 1.0 if 3.5 <= avg_len <= 8 else max(0, 1 - abs(avg_len - 5.5) / 5)

    # Common English stop-words present?
    stops = {'the', 'and', 'in', 'of', 'to', 'a', 'is', 'for', 'with', 'on', 'at', 'an', 'or'}
    lower_words = {w.lower().strip('.,;:') for w in words}
    stop_hit = len(stops & lower_words) / len(stops)

    # Penalise excessive non-alnum chars (bullet-stuffed text)
    alnum_count = sum(1 for c in text if c.isalnum())
    alnum_ratio = alnum_count / len(text) if text else 0
    # Good text: ~70-85% alnum. Below 40% is suspicious.
    alnum_penalty = max(0, min(1, (alnum_ratio - 0.2) / 0.4))  # 0 at <=20%, 1 at >=60%

    score = (1 - single_char_ratio) * 0.3 + length_score * 0.2 + stop_hit * 0.2 + alnum_penalty * 0.3
    return max(0.0, min(1.0, score))

def try_fix_encoding(text: str) -> str:
    """Attempt to repair double-encoded UTF-8 text."""
    if not text:
        return text
    try:
        fixed = text.encode('cp1252').decode('utf-8', errors='ignore')
        if fixed and len(fixed) > len(text) * 0.3 and not is_mojibake(fixed, 0.1):
            return fixed
    except (UnicodeEncodeError, UnicodeDecodeError):
        pass
    try:
        fixed = text.encode('latin-1').decode('utf-8', errors='ignore')
        if fixed and len(fixed) > len(text) * 0.3 and not is_mojibake(fixed, 0.1):
            return fixed
    except (UnicodeEncodeError, UnicodeDecodeError):
        pass
    return text


class ResumeParser:
    """
    Service for parsing resume files and extracting structured data.
    
    Uses a tiered approach:
    1. LLM (Ollama) for 100% accurate structured extraction (primary)
    2. Regex-based extraction as fallback when LLM is unavailable
    """
    
    def __init__(self):
        self._llm_service = None
        # Comprehensive skill keywords for better extraction
        self.skill_keywords = [
            # Programming Languages
            'python', 'javascript', 'typescript', 'react', 'node.js', 'vue', 'angular',
            'java', 'c++', 'c#', '.net', 'go', 'golang', 'rust', 'ruby', 'php', 'swift',
            'kotlin', 'scala', 'perl', 'matlab', 'r programming', 'dart', 'flutter',
            
            # Databases
            'sql', 'postgresql', 'mongodb', 'mysql', 'redis', 'elasticsearch', 'oracle',
            'sqlite', 'cassandra', 'dynamodb', 'firebase', 'neo4j', 'mariadb',
            
            # DevOps & Cloud
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible',
            'jenkins', 'ci/cd', 'github actions', 'gitlab', 'bitbucket', 'circleci',
            'prometheus', 'grafana', 'nginx', 'apache', 'linux', 'unix', 'bash',
            
            # Frameworks & Tools
            'django', 'flask', 'fastapi', 'express', 'spring boot', 'hibernate',
            'rails', 'laravel', 'nextjs', 'nuxt', 'gatsby', 'svelte', 'redux',
            'tailwind', 'bootstrap', 'material ui', 'sass', 'webpack', 'vite',
            
            # AI/ML & Data
            'machine learning', 'deep learning', 'ai', 'data science', 'tensorflow',
            'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'spark', 'hadoop',
            'tableau', 'power bi', 'jupyter', 'nlp', 'computer vision', 'opencv',
            
            # Soft Skills & Methodologies
            'agile', 'scrum', 'kanban', 'jira', 'confluence', 'trello', 'asana',
            'leadership', 'management', 'communication', 'teamwork', 'problem solving',
            
            # Architecture & Concepts
            'rest', 'graphql', 'microservices', 'api', 'backend', 'frontend',
            'full-stack', 'devops', 'cloud', 'saas', 'paas', 'serverless',
            'event-driven', 'message queue', 'rabbitmq', 'kafka', 'websocket',
            
            # Security
            'security', 'oauth', 'jwt', 'encryption', 'ssl', 'penetration testing',
            'vulnerability assessment', 'firewall', 'siem',
            
            # Business & Marketing
            'seo', 'google analytics', 'marketing', 'sales', 'crm', 'salesforce',
            'hubspot', 'mailchimp', 'social media', 'content marketing', 'ppc',
            
            # Design
            'figma', 'sketch', 'adobe xd', 'photoshop', 'illustrator', 'ui/ux',
            'user research', 'wireframing', 'prototyping', 'design systems',
            
            # Mobile
            'ios', 'android', 'react native', 'xamarin', 'ionic',
            
            # Version Control
            'git', 'github', 'version control', 'code review'
        ]
    
    async def extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX file"""
        if filename.lower().endswith('.pdf'):
            return self._extract_from_pdf(content)
        elif filename.lower().endswith('.docx'):
            return self._extract_from_docx(content)
        else:
            raise ValueError("Unsupported file format. Supported: PDF, DOCX")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted PDF text for better parsing accuracy"""
        if not text:
            return ""

        # Step 0: Detect and repair spaced-character corruption
        if is_spaced_text(text):
            logger.warning("Detected spaced-character corruption - collapsing")
            text = collapse_spaced_chars(text)

        # Try to fix double-encoded UTF-8 (mojibake)
        if is_mojibake(text):
            text = try_fix_encoding(text)

        # Unicode normalization
        text = unicodedata.normalize('NFC', text)

        # Strip CID font artifacts: (cid:NN) sequences from broken PDF font mappings
        # These appear when PDFs use CID (Character ID) fonts with custom encodings
        if '(cid:' in text:
            cid_count = len(re.findall(r'\(cid:\d+\)', text))
            if cid_count > 0:
                logger.warning(f"Stripping {cid_count} CID font artifacts from PDF text")
                text = re.sub(r'\(cid:\d+\)', '', text)
                # Also strip orphaned parentheses left from partial CID patterns
                text = re.sub(r'\(cid:[^)]*\)', '', text)

        # Fix common PDF extraction artifacts
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Collapse multiple spaces (but preserve intentional indentation)
            line = re.sub(r'  +', '  ', line)
            # Remove control characters except newline/tab
            line = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', line)
            # Fix broken words from line wrapping (e.g., "Soft- ware" -> "Software")
            # Only join if both sides are >=2 chars to avoid mangling spaced-char leftovers
            line = re.sub(r'(\w{2,})-\s+(\w{2,})', r'\1\2', line)
            cleaned_lines.append(line.strip())

        text = '\n'.join(cleaned_lines)

        # Remove excessive blank lines (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Normalize dashes
        text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2015', '-')

        # Normalize quotes
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")

        return text.strip()

    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using multiple methods with quality checks.
        
        Strategy:
        1. Try pdfplumber layout-aware extraction
        2. If quality is poor (spaced chars), try pdfplumber basic mode
        3. If still poor, fall back to PyPDF2
        4. Pick the best quality result
        """
        pdf_file = BytesIO(content)
        candidates = []  # (quality_score, text, method_name)

        # --- Method 1: pdfplumber layout-aware ---
        try:
            import pdfplumber
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text(
                        layout=True,
                        x_density=7.25,
                        y_density=13
                    )
                    if page_text:
                        text += page_text + "\n\n"
                    # Also extract text from tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = ' | '.join([cell or '' for cell in row])
                                    if row_text.strip('| '):
                                        text += row_text + "\n"

            if text.strip():
                cleaned = self._clean_extracted_text(text)
                score = text_quality_score(cleaned)
                candidates.append((score, cleaned, "pdfplumber-layout"))
                logger.info(f"pdfplumber layout: {len(cleaned)} chars, quality={score:.2f}")
                if score >= 0.5:
                    # Good enough, return immediately
                    return cleaned
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pdfplumber layout extraction failed: {e}")

        # --- Method 2: pdfplumber basic (no layout) ---
        try:
            pdf_file.seek(0)
            import pdfplumber
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            if text.strip():
                cleaned = self._clean_extracted_text(text)
                score = text_quality_score(cleaned)
                candidates.append((score, cleaned, "pdfplumber-basic"))
                logger.info(f"pdfplumber basic: {len(cleaned)} chars, quality={score:.2f}")
                if score >= 0.5:
                    return cleaned
        except Exception as e:
            logger.debug(f"pdfplumber basic extraction failed: {e}")

        # --- Method 3: PyPDF2 fallback ---
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            if text.strip():
                cleaned = self._clean_extracted_text(text)
                score = text_quality_score(cleaned)
                candidates.append((score, cleaned, "PyPDF2"))
                logger.info(f"PyPDF2: {len(cleaned)} chars, quality={score:.2f}")
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed: {e}")

        # Pick the best candidate by quality score
        if candidates:
            candidates.sort(key=lambda c: c[0], reverse=True)
            best_score, best_text, best_method = candidates[0]
            logger.info(f"Selected {best_method} (quality={best_score:.2f}, {len(best_text)} chars)")
            return best_text

        logger.warning("All PDF extraction methods failed")
        return ""

    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX including paragraphs and tables"""
        doc_file = BytesIO(content)
        doc = docx.Document(doc_file)
        
        # Extract paragraph text
        parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract table text (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        
        return "\n".join(parts)
    
    async def parse_resume(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse resume and extract structured data.
        Uses LLM for 100% accurate extraction, falls back to regex if unavailable.
        """
        text = await self.extract_text(content, filename)
        
        if not text or len(text.strip()) < 30:
            logger.warning(f"Insufficient text extracted from {filename}")
            return self._empty_result(text)
        
        # Check for mojibake/garbled text — flag it if mostly unreadable
        if is_mojibake(text, threshold=0.25):
            logger.warning(f"Text from {filename} appears garbled/mojibake ({len(text)} chars). Attempting repair.")
            text = try_fix_encoding(text)
            if is_mojibake(text, threshold=0.25):
                logger.warning(f"Text still garbled after repair for {filename}")
        
        # Strategy 1: Try LLM-powered extraction (100% accurate)
        llm_result = await self._parse_with_llm(text)
        if llm_result:
            logger.info(f"Resume parsed with LLM: {llm_result.get('name', 'Unknown')}")
            llm_result['raw_text'] = text[:8000]
            return llm_result
        
        # Strategy 2: Try Gemini-powered extraction (production fallback)
        gemini_result = await self._parse_with_gemini(text)
        if gemini_result:
            logger.info(f"Resume parsed with Gemini: {gemini_result.get('name', 'Unknown')}")
            gemini_result['raw_text'] = text[:8000]
            return gemini_result
        
        # Strategy 3: Fallback to regex-based extraction
        logger.info("LLM+Gemini unavailable, using regex-based extraction")
        name = self._extract_name(text)
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        skills = self._extract_skills(text)
        experience = self._extract_experience(text)
        education = self._extract_education(text)
        work_history = self._extract_work_history(text)
        summary = self._extract_summary(text)
        location = self._extract_location(text)
        linkedin = self._extract_linkedin(text)
        
        return {
            "name": name,
            "email": email,
            "phone": phone,
            "skills": skills,
            "experience": experience,
            "education": education,
            "work_history": work_history,
            "summary": summary,
            "location": location,
            "linkedin": linkedin,
            "certifications": [],
            "languages": [],
            "raw_text": text[:5000],
            "parsed_by": "regex"
        }
    
    async def _parse_with_llm(self, text: str) -> Optional[Dict[str, Any]]:
        """Parse resume using LLM service for 100% accurate extraction"""
        try:
            # Lazy import to avoid circular dependencies
            if self._llm_service is None:
                from services.llm_service import get_llm_service
                self._llm_service = await get_llm_service()
            
            if not self._llm_service.available:
                return None
            
            result = await self._llm_service.parse_resume(text)
            
            if result and (result.get('name') or result.get('email') or result.get('skills')):
                # Convert LLM format to parser format
                return {
                    "name": result.get('name', 'Unknown'),
                    "email": result.get('email', ''),
                    "phone": result.get('phone', ''),
                    "skills": result.get('skills', []),
                    "experience": result.get('experience_years', 0),
                    "education": result.get('education', []),
                    "work_history": result.get('work_history', []),
                    "summary": result.get('summary', ''),
                    "location": result.get('location', ''),
                    "linkedin": result.get('linkedin', ''),
                    "certifications": result.get('certifications', []),
                    "languages": result.get('languages', []),
                    "job_category": result.get('job_category', 'General'),
                    "parsed_by": "llm"
                }
            
            return None
            
        except Exception as e:
            logger.warning(f"LLM resume parsing failed: {e}")
            return None
    
    async def _parse_with_gemini(self, text: str) -> Optional[Dict[str, Any]]:
        """Parse resume using Gemini AI service as production fallback.
        
        This is critical for Cloud Run where Ollama isn't available.
        Uses gemini.parse_resume() which has a dedicated extraction prompt
        with comprehensive field extraction.
        """
        try:
            from services.gemini_service import get_gemini_service
            gemini = get_gemini_service()
            if not gemini or not gemini.available:
                logger.info("Gemini service not available for resume parsing")
                return None
            
            # Use parse_resume which has a dedicated extraction prompt
            # (NOT analyze_candidate which is for scoring/quality assessment)
            result = await gemini.parse_resume(text)
            
            if result and (result.get('name') or result.get('email') or result.get('skills')):
                # Normalize work_history format
                work_history = result.get('work_history', [])
                if isinstance(work_history, list):
                    normalized_wh = []
                    for w in work_history:
                        if isinstance(w, dict):
                            normalized_wh.append(w)
                        elif isinstance(w, str):
                            normalized_wh.append({"title": w, "company": "", "period": "", "description": ""})
                    work_history = normalized_wh
                
                # Normalize education format
                education = result.get('education', [])
                if isinstance(education, list):
                    normalized_edu = []
                    for e in education:
                        if isinstance(e, dict):
                            normalized_edu.append(e)
                        elif isinstance(e, str):
                            normalized_edu.append({"degree": e, "field": "", "institution": "", "year": ""})
                    education = normalized_edu
                
                return {
                    "name": result.get('name', 'Unknown'),
                    "email": result.get('email', ''),
                    "phone": result.get('phone', ''),
                    "skills": result.get('skills', []),
                    "experience": result.get('experience', result.get('experience_years', 0)),
                    "education": education,
                    "work_history": work_history,
                    "summary": result.get('summary', ''),
                    "location": result.get('location', ''),
                    "linkedin": result.get('linkedin', ''),
                    "certifications": result.get('certifications', []),
                    "languages": result.get('languages', []),
                    "job_category": result.get('job_category', 'General'),
                    "job_subcategory": result.get('job_subcategory', ''),
                    "nationality": result.get('nationality', ''),
                    "parsed_by": "gemini"
                }
            
            return None
            
        except Exception as e:
            logger.warning(f"Gemini resume parsing failed: {e}")
            return None
    
    def _empty_result(self, text: str = "") -> Dict[str, Any]:
        """Return empty result structure"""
        return {
            "name": "Unknown",
            "email": "",
            "phone": "",
            "skills": [],
            "experience": 0,
            "education": [],
            "work_history": [],
            "summary": "",
            "location": "",
            "linkedin": "",
            "certifications": [],
            "languages": [],
            "raw_text": text[:5000] if text else "",
            "parsed_by": "none"
        }
    
    def _extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn profile URL"""
        # LinkedIn URL patterns
        patterns = [
            r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+/?',
            r'linkedin\.com/in/[\w\-]+',
            r'linkedin:\s*([\w\-]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                url = matches[0]
                # Ensure it's a full URL
                if not url.startswith('http'):
                    if url.startswith('linkedin.com'):
                        url = 'https://' + url
                    else:
                        url = f'https://linkedin.com/in/{url}'
                return url
        
        return ""
    
    def _extract_location(self, text: str) -> str:
        """Extract location from resume using patterns and city/country matching"""
        # First try structured patterns: "Location: City, Country" or "Address: ..."
        loc_patterns = [
            r'(?:location|address|city|based in|residing in)[:\s]+([A-Za-z\s,]+(?:,\s*[A-Za-z\s]+))',
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*(?:UAE|United Arab Emirates|India|Pakistan|USA|UK|Canada|Australia|Singapore|Philippines|Qatar|Saudi Arabia|Oman|Bahrain|Kuwait|Jordan|Egypt|Lebanon|Germany|France|Netherlands|Ireland))',
        ]
        for pattern in loc_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                loc = match.group(1).strip().rstrip(',')
                if len(loc) > 2 and len(loc) < 60:
                    return loc.title()
        
        text_lower = text.lower()
        
        # UAE cities (specific first)
        uae_cities = {
            'dubai': 'Dubai, UAE', 'abu dhabi': 'Abu Dhabi, UAE', 'sharjah': 'Sharjah, UAE',
            'ajman': 'Ajman, UAE', 'fujairah': 'Fujairah, UAE',
            'ras al khaimah': 'Ras Al Khaimah, UAE', 'umm al quwain': 'Umm Al Quwain, UAE'
        }
        for city, full_loc in uae_cities.items():
            if city in text_lower:
                return full_loc
        
        # General country/region matching
        country_map = {
            'united arab emirates': 'UAE', 'uae': 'UAE',
            'india': 'India', 'pakistan': 'Pakistan', 'philippines': 'Philippines',
            'united kingdom': 'UK', 'uk': 'UK', 'united states': 'USA', 'usa': 'USA',
            'canada': 'Canada', 'australia': 'Australia', 'singapore': 'Singapore',
            'qatar': 'Qatar', 'saudi arabia': 'Saudi Arabia', 'oman': 'Oman',
            'bahrain': 'Bahrain', 'kuwait': 'Kuwait', 'jordan': 'Jordan',
            'egypt': 'Egypt', 'lebanon': 'Lebanon', 'germany': 'Germany',
            'france': 'France', 'netherlands': 'Netherlands', 'ireland': 'Ireland',
            'new zealand': 'New Zealand', 'south africa': 'South Africa',
            'remote': 'Remote', 'hong kong': 'Hong Kong', 'malaysia': 'Malaysia',
            'indonesia': 'Indonesia', 'thailand': 'Thailand', 'vietnam': 'Vietnam',
            'china': 'China', 'japan': 'Japan', 'south korea': 'South Korea',
            'brazil': 'Brazil', 'nigeria': 'Nigeria', 'kenya': 'Kenya',
            'sri lanka': 'Sri Lanka', 'bangladesh': 'Bangladesh', 'nepal': 'Nepal',
        }
        for keyword, country in country_map.items():
            if keyword in text_lower:
                return country
        
        return ''
    
    def _is_valid_name(self, name: str) -> bool:
        """Check if extracted name is valid (not garbage)"""
        if not name or len(name) < 2 or len(name) > 60:
            return False
        # Contains too many digits (likely a date/ID)
        digit_count = sum(1 for c in name if c.isdigit())
        if digit_count > 3:
            return False
        # Contains date patterns
        if re.search(r'\d{4}\s*[-â€“]\s*\d{4}', name.replace(' ', '')):
            return False
        if re.search(r'[0-9]\s+[0-9]\s+[0-9]\s+[0-9]', name):
            return False
        # Is a common keyword
        garbage_keywords = ['resume', 'cv', 'curriculum', 'vitae', 'summary', 'professional', 
                          'objective', 'experience', 'education', 'skills', 'contact', 
                          'personal', 'details', 'career', 'profile']
        name_lower = name.lower().strip()
        for kw in garbage_keywords:
            if name_lower == kw or name_lower.startswith(kw + ' ') or name_lower.endswith(' ' + kw):
                return False
        # At least one letter
        if not any(c.isalpha() for c in name):
            return False
        return True
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume text - improved algorithm"""
        lines = text.split('\n')
        
        # First pass: look for lines that look like names in the first 10 lines
        for line in lines[:10]:
            line = line.strip()
            # Skip empty, emails, or very long lines
            if not line or '@' in line or len(line) > 60:
                continue
            # Skip lines with too many digits
            if sum(1 for c in line if c.isdigit()) > 2:
                continue
            # Skip lines that are clearly headers
            if any(kw in line.lower() for kw in ['resume', 'cv', 'curriculum', 'summary', 'experience', 'education', 'skills', 'objective', 'contact', 'profile']):
                continue
            # Valid name: 2-4 words, mostly letters
            words = line.split()
            if 1 <= len(words) <= 5 and len(line) >= 3 and len(line) <= 50:
                if self._is_valid_name(line):
                    return line
        
        # Fallback: try more relaxed matching
        for line in lines[:15]:
            line = line.strip()
            if 3 <= len(line) <= 40 and '@' not in line:
                if self._is_valid_name(line):
                    return line
        
        return "Unknown"
    
    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else ""
    
    @staticmethod
    def _is_valid_phone(phone: str) -> bool:
        """Validate a phone string is actually a phone number, not CID/garbage."""
        if not phone:
            return False
        # Reject CID font artifacts
        if 'cid:' in phone.lower() or '(cid' in phone.lower():
            return False
        # Must contain at least 7 digits
        digits = re.sub(r'\D', '', phone)
        if len(digits) < 7 or len(digits) > 15:
            return False
        # Reject if more than 50% non-digit, non-standard chars
        allowed = set('0123456789+()-. ')
        if sum(1 for c in phone if c not in allowed) > len(phone) * 0.3:
            return False
        return True

    def _extract_phone(self, text: str) -> str:
        """Extract phone number (supports UAE +971 and international formats)"""
        # Strip CID artifacts before phone extraction
        clean_text = re.sub(r'\(cid:\d+\)', '', text)
        
        # UAE format: +971-XX-XXX-XXXX or variations
        uae_pattern = r'\+971[\s.-]?\d{1,2}[\s.-]?\d{3}[\s.-]?\d{4}'
        uae_matches = re.findall(uae_pattern, clean_text)
        if uae_matches and self._is_valid_phone(uae_matches[0]):
            return uae_matches[0]
        
        # India format: +91-XXXXX-XXXXX
        india_pattern = r'\+91[\s.-]?\d{5}[\s.-]?\d{5}'
        india_matches = re.findall(india_pattern, clean_text)
        if india_matches and self._is_valid_phone(india_matches[0]):
            return india_matches[0]
        
        # General international format: +CC-XXX-XXX-XXXX
        intl_pattern = r'\+\d{1,3}[\s.-]?\(?\d{2,4}\)?[\s.-]?\d{3,5}[\s.-]?\d{3,5}'
        intl_matches = re.findall(intl_pattern, clean_text)
        if intl_matches and self._is_valid_phone(intl_matches[0]):
            return intl_matches[0]
        
        # US format: (XXX) XXX-XXXX
        us_pattern = r'\(\d{3}\)\s*\d{3}[\s.-]?\d{4}'
        us_matches = re.findall(us_pattern, clean_text)
        if us_matches and self._is_valid_phone(us_matches[0]):
            return us_matches[0]
        
        # Plain digit sequences (10-12 digits)
        plain_pattern = r'\b\d{10,12}\b'
        plain_matches = re.findall(plain_pattern, clean_text)
        if plain_matches and self._is_valid_phone(plain_matches[0]):
            return plain_matches[0]
        
        # General fallback — but validate
        phone_pattern = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
        matches = re.findall(phone_pattern, clean_text)
        for m in matches:
            if self._is_valid_phone(m):
                return m
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills using word-boundary matching"""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.skill_keywords:
            # Use word boundaries to prevent false positives
            # e.g., 'go' shouldn't match 'going', 'r' shouldn't match 'required'
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found_skills.append(skill.title())
        
        return list(set(found_skills))[:30]  # Return up to 30 unique skills
    
    def _extract_experience(self, text: str) -> int:
        """Extract years of experience"""
        text_lower = text.lower()
        # Look for patterns like "5 years", "5+ years", etc.
        patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+in\s+',
            r'(\d+)\+?\s*yrs?\s+(?:of\s+)?exp',
            r'(\d+)\+?\s*years?\s+(?:working|professional)',
            r'over\s+(\d+)\+?\s*years?',
        ]
        
        max_exp = 0
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    val = int(match)
                    if 0 < val < 50:  # Reasonable range
                        max_exp = max(max_exp, val)
                except Exception:
                    pass
        
        return max_exp
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information with improved patterns and year extraction"""
        education = []
        original_text = text
        
        # Helper to find graduation year near a match position
        def find_year_near(pos: int, search_text: str) -> str:
            # Search within 150 chars around the match for a 4-digit year
            context = search_text[max(0, pos - 80):pos + 150]
            year_match = re.search(r'(19[89]\d|20[0-2]\d)', context)
            return year_match.group(1) if year_match else ''
        
        # Comprehensive degree patterns â€” don't break on first match, collect all
        full_edu_patterns = [
            # PhD patterns
            (r'(?:ph\.?d\.?|doctorate?|doctor of philosophy)\s+(?:in\s+|of\s+)?([A-Za-z\s&]{3,40})\s+(?:from\s+)?([A-Za-z\s]+(?:University|College|Institute|School))', 'PhD'),
            (r'(?:ph\.?d\.?|doctorate?)\s+(?:in\s+)?([A-Za-z\s&]{3,30})(?:\s|,|$)', 'PhD'),
            # Masters patterns
            (r"(?:master'?s?|m\.?s\.?c?\.?|mba|m\.?a\.?|m\.?tech|m\.?eng|m\.?phil)\s+(?:in\s+|of\s+)?([A-Za-z\s&]{3,40})\s+(?:from\s+)?([A-Za-z\s]+(?:University|College|Institute|School))", 'Masters'),
            (r"(?:master'?s?\s+(?:of\s+|in\s+)?(?:science|arts|business|engineering)|m\.?s\.?c?\.?|mba|m\.?tech)\s+(?:in\s+)?([A-Za-z\s&]{3,30})", 'Masters'),
            # Bachelors patterns
            (r"(?:bachelor'?s?|b\.?s\.?c?\.?|b\.?a\.?|b\.?e\.?|b\.?tech|b\.?eng|b\.?com|b\.?b\.?a)\s+(?:in\s+|of\s+)?([A-Za-z\s&]{3,40})\s+(?:from\s+)?([A-Za-z\s]+(?:University|College|Institute|School))", 'Bachelors'),
            (r"(?:bachelor'?s?\s+(?:of\s+|in\s+)?(?:science|arts|engineering|technology|commerce)|b\.?e\.?|b\.?tech|b\.?s\.?c?\.?|b\.?com|b\.?b\.?a)\s+(?:in\s+)?([A-Za-z\s&]{3,30})", 'Bachelors'),
            # Associate/Diploma patterns
            (r"(?:associate'?s?|diploma|a\.?s\.?|a\.?a\.?)\s+(?:in\s+|of\s+)?([A-Za-z\s&]{3,40})", 'Diploma'),
        ]
        
        seen_degrees = set()
        for pattern, degree_type in full_edu_patterns:
            for match in re.finditer(pattern, original_text, re.IGNORECASE):
                field = match.group(1).strip() if match.group(1) else ''
                institution = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ''
                # Clean up field
                field = re.sub(r'\b(in|of|from|the|and|with|at)\b', '', field, flags=re.IGNORECASE).strip()
                field = field.strip(' ,.-')
                
                if len(field) >= 3 and not re.match(r'^[a-z]{1,3}$', field.lower()):
                    # Deduplicate
                    degree_key = f"{degree_type}_{field.lower()}"
                    if degree_key not in seen_degrees:
                        seen_degrees.add(degree_key)
                        year = find_year_near(match.start(), original_text)
                        education.append({
                            "degree": degree_type,
                            "field": field.title(),
                            "institution": institution.title() if institution else '',
                            "year": year
                        })
        
        # Fallback: look for university mentions
        if not education:
            for uni_match in re.finditer(r'([A-Z][a-zA-Z\s]+(?:University|College|Institute|School|Academy))', original_text):
                # Check for degree keywords nearby
                text_around = original_text[max(0, uni_match.start()-120):uni_match.end()+60].lower()
                degree = 'Degree'
                field = ''
                if 'phd' in text_around or 'doctor' in text_around:
                    degree = 'PhD'
                elif 'master' in text_around or 'mba' in text_around or 'm.s' in text_around:
                    degree = 'Masters'
                elif 'bachelor' in text_around or 'b.e' in text_around or 'b.tech' in text_around or 'b.sc' in text_around:
                    degree = 'Bachelors'
                elif 'diploma' in text_around or 'associate' in text_around:
                    degree = 'Diploma'
                
                year = find_year_near(uni_match.start(), original_text)
                institution = uni_match.group(1).strip()
                
                # Deduplicate
                if institution not in [e.get('institution', '') for e in education]:
                    education.append({
                        "degree": degree,
                        "field": field,
                        "institution": institution,
                        "year": year
                    })
                    if len(education) >= 3:
                        break
        
        return education[:5]
    
    def _extract_work_history(self, text: str) -> List[Dict[str, str]]:
        """Extract work history from resume with company detection"""
        work_history = []
        
        # Common job title keywords
        job_titles = [
            'engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant',
            'specialist', 'coordinator', 'director', 'executive', 'lead', 'architect',
            'administrator', 'associate', 'assistant', 'intern', 'trainee', 'officer',
            'supervisor', 'head of', 'vice president', 'vp', 'cto', 'ceo', 'cfo',
            'founder', 'co-founder', 'partner', 'accountant', 'recruiter', 'advisor',
            'strategist', 'planner', 'controller', 'auditor', 'scientist', 'researcher'
        ]
        
        # Company indicator words
        company_indicators = ['at', 'for', '@', '-', '|', 'Â·', ',']
        
        lines = text.split('\n')
        i = 0
        while i < len(lines) and len(work_history) < 8:
            line = lines[i].strip()
            line_lower = line.lower()
            
            # Check if this line contains a job title
            has_title = any(title in line_lower for title in job_titles)
            
            if has_title and 5 < len(line) < 200:
                # Extract year range from this line or the next
                year_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?\s*(20\d{2}|19\d{2})\s*[-â€“to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?\s*(20\d{2}|[Pp]resent|[Cc]urrent|[Oo]ngoing)', line, re.IGNORECASE)
                period = year_match.group(0).strip() if year_match else ''
                
                # If no period found on this line, check next line
                if not period and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    year_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?\s*(20\d{2}|19\d{2})\s*[-â€“to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?\s*(20\d{2}|[Pp]resent|[Cc]urrent|[Oo]ngoing)', next_line, re.IGNORECASE)
                    if year_match:
                        period = year_match.group(0).strip()
                
                # Try to separate title from company
                title_text = line
                company = ''
                
                # Pattern: "Title at/@ Company" or "Title - Company" or "Title | Company"
                for sep in [' at ', ' @ ', ' - ', ' | ', ' Â· ']:
                    if sep in line:
                        parts = line.split(sep, 1)
                        # The part with the job title keyword is the title
                        if any(t in parts[0].lower() for t in job_titles):
                            title_text = parts[0].strip()
                            company = parts[1].strip()
                        else:
                            company = parts[0].strip()
                            title_text = parts[1].strip()
                        break
                
                # Also check if next/nearby line is the company (often on separate line)
                if not company and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    # If next line is short and doesn't contain job title keywords, it might be company
                    if (5 < len(next_line) < 80 and
                        not any(t in next_line.lower() for t in job_titles) and
                        not next_line.startswith(('â€¢', '-', '*', 'Â·'))):
                        # Check if it looks like a company name (starts with capital, no bullets)
                        if next_line[0].isupper() or next_line[0].isdigit():
                            company = next_line
                
                # Remove period/dates from title
                if period and period in title_text:
                    title_text = title_text.replace(period, '').strip().rstrip('-|Â·,')
                if period and company and period in company:
                    company = company.replace(period, '').strip().rstrip('-|Â·,')
                
                # Collect description from bullet points following this entry
                description_lines = []
                j = i + 1
                while j < min(i + 10, len(lines)):
                    desc_line = lines[j].strip()
                    if desc_line.startswith(('â€¢', '-', '*', 'Â·', 'â—‹')):
                        clean_desc = desc_line.lstrip('â€¢-*Â·â—‹ ')
                        if len(clean_desc) > 10:
                            description_lines.append(clean_desc)
                    elif len(desc_line) > 0 and any(t in desc_line.lower() for t in job_titles):
                        break  # Next job entry
                    j += 1
                
                description = '; '.join(description_lines[:4]) if description_lines else ''
                
                work_history.append({
                    "title": title_text[:100],
                    "company": company[:80],
                    "period": period,
                    "description": description[:500]
                })
            i += 1
        
        return work_history[:8]
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary with better section detection"""
        summary_keywords = ['professional summary', 'summary', 'profile', 'about me',
                           'about', 'overview', 'objective', 'career summary',
                           'executive summary', 'personal statement', 'career objective']
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_stripped = line.strip().lower().rstrip(':')
            # Check if line IS a header (short line that matches a keyword)
            if any(keyword == line_stripped or line_stripped.endswith(keyword) for keyword in summary_keywords):
                # Collect subsequent lines until next section header or blank gap
                summary_lines = []
                for j in range(i + 1, min(i + 12, len(lines))):
                    next_line = lines[j].strip()
                    if not next_line:
                        if summary_lines:  # Stop at blank line after content
                            break
                        continue
                    # Stop if we hit another section header
                    if next_line.lower().rstrip(':') in ['experience', 'education', 'skills',
                        'work experience', 'employment', 'certifications', 'technical skills',
                        'professional experience', 'work history', 'projects']:
                        break
                    summary_lines.append(next_line)
                
                summary = ' '.join(summary_lines).strip()
                if len(summary) > 30:
                    return summary[:600]
        
        # Fallback: find first substantial paragraph (after name/contact info)
        paragraphs = text.split('\n\n')
        for para in paragraphs[1:]:  # Skip first block (likely name/contact)
            para = para.strip()
            # Must be substantive but not a list of skills
            if len(para) > 80 and not para.startswith(('â€¢', '-', '*')):
                # Skip if it looks like a header + bullet list
                lines_in_para = para.split('\n')
                non_bullet = [l for l in lines_in_para if not l.strip().startswith(('â€¢', '-', '*'))]
                text_portion = ' '.join(non_bullet).strip()
                if len(text_portion) > 60:
                    return text_portion[:600]
        
        return ''
    
    async def parse_job_description(self, text: str) -> Dict[str, Any]:
        """Parse job description and extract requirements"""
        
        required_skills = self._extract_skills(text)
        preferred_skills = []  # Extract preferred skills separately
        experience_level = self._extract_experience_level(text)
        responsibilities = self._extract_responsibilities(text)
        
        return {
            "title": self._extract_job_title(text),
            "required_skills": required_skills[:10],
            "preferred_skills": preferred_skills,
            "experience_level": experience_level,
            "responsibilities": responsibilities,
            "description": text[:500]
        }
    
    def _extract_job_title(self, text: str) -> str:
        """Extract job title"""
        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 5 and len(line) < 80:
                return line
        return "Software Engineer"
    
    def _extract_experience_level(self, text: str) -> str:
        """Extract experience level"""
        text_lower = text.lower()
        
        if 'senior' in text_lower or '5+ years' in text_lower:
            return "Senior (5+ years)"
        elif 'junior' in text_lower or 'entry' in text_lower:
            return "Junior (0-2 years)"
        else:
            return "Mid-level (2-5 years)"
    
    def _extract_responsibilities(self, text: str) -> List[str]:
        """Extract key responsibilities"""
        responsibilities = []
        lines = text.split('\n')
        
        in_responsibilities = False
        for line in lines:
            line = line.strip()
            
            if 'responsibilities' in line.lower() or 'duties' in line.lower():
                in_responsibilities = True
                continue
            
            if in_responsibilities and line.startswith(('-', 'â€¢', '*')):
                resp = line.lstrip('-â€¢* ').strip()
                if len(resp) > 20:
                    responsibilities.append(resp)
                    if len(responsibilities) >= 5:
                        break
        
        return responsibilities
